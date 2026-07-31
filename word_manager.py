from docx import Document
from docx.shared import Pt
from copy import deepcopy
import os

from config import CURRENT_WORD, TEMP_FOLDER


class WordManager:

    def __init__(self):
        self.document = None
        self.table = None

    def abrir(self):
        self.document = Document(CURRENT_WORD)

        for tabela in self.document.tables:
            if len(tabela.rows) > 20:
                self.table = tabela
                break

        if self.table is None:
            raise Exception("Tabela da folha de ponto não encontrada.")

    def salvar(self, nome_arquivo):
        os.makedirs(TEMP_FOLDER, exist_ok=True)

        caminho = os.path.join(
            TEMP_FOLDER,
            nome_arquivo
        )

        self.document.save(caminho)

        return caminho

    def localizar_linha(self, dia):

        for row in self.table.rows:

            texto = row.cells[0].text.strip()

            if texto == f"{dia:02d}":
                return row

        return None
    def escrever(self, celula, texto):
        """
        Escreve um texto na célula preservando a formatação básica.
        """
        celula.text = str(texto)

        for paragrafo in celula.paragraphs:
            for run in paragrafo.runs:
                run.font.size = Pt(10)

    def preencher_dia(
        self,
        dia,
        entrada="",
        intervalo_inicio="",
        intervalo_fim="",
        saida=""
    ):
        """
        Preenche os horários de um determinado dia.
        """

        linha = self.localizar_linha(dia)

        if linha is None:
            raise Exception(f"Dia {dia:02d} não encontrado na folha.")

        # Colunas da tabela
        # 2 = Entrada
        # 3 = Início intervalo
        # 4 = Fim intervalo
        # 5 = Saída

        self.escrever(linha.cells[2], entrada)
        self.escrever(linha.cells[3], intervalo_inicio)
        self.escrever(linha.cells[4], intervalo_fim)
        self.escrever(linha.cells[5], saida)

    def limpar_dia(self, dia):
        linha = self.localizar_linha(dia)

        if linha is None:
            return

        for coluna in range(2, 6):
            self.escrever(linha.cells[coluna], "")    def preencher_varios_dias(self, registros):
        """
        Preenche vários dias de uma só vez.

        registros = {
            1: {
                "entrada": "08:00",
                "intervalo_inicio": "12:00",
                "intervalo_fim": "13:00",
                "saida": "17:30"
            },
            2: {
                ...
            }
        }
        """

        for dia, dados in registros.items():

            self.preencher_dia(
                int(dia),
                dados.get("entrada", ""),
                dados.get("intervalo_inicio", ""),
                dados.get("intervalo_fim", ""),
                dados.get("saida", "")
            )

    def carregar_modelo(self):
        """
        Carrega novamente o modelo original da folha.
        """

        self.abrir()

    def gerar_folha(self, registros, nome_saida):
        """
        Gera uma folha preenchida.
        """

        self.carregar_modelo()

        self.preencher_varios_dias(registros)

        return self.salvar(nome_saida)
